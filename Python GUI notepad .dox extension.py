import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
from docx.shared import Pt

class NotepadApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Simple Notepad App")
        
        self.create_widgets()
        
    def create_widgets(self):
        self.textbox = tk.Text(self.root, wrap="word")
        self.textbox.pack(fill="both", expand=True)
        
        self.format_button = ttk.Button(self.root, text="Format Text", command=self.open_format_window)
        self.format_button.pack(side="left", padx=5, pady=5)
        
        self.save_button = ttk.Button(self.root, text="Save", command=self.save_document)
        self.save_button.pack(side="left", padx=5, pady=5)
        
        self.open_button = ttk.Button(self.root, text="Open", command=self.open_document)
        self.open_button.pack(side="left", padx=5, pady=5)
        
        self.append_button = ttk.Button(self.root, text="Append", command=self.append_to_document)
        self.append_button.pack(side="left", padx=5, pady=5)
        
    def open_format_window(self):
        format_window = tk.Toplevel(self.root)
        format_window.title("Text Formatting")
        
        bold_button = ttk.Button(format_window, text="Bold", command=lambda: self.apply_format("<b>"))
        bold_button.grid(row=0, column=0, padx=5, pady=5)
        
        italic_button = ttk.Button(format_window, text="Italic", command=lambda: self.apply_format("<i>"))
        italic_button.grid(row=0, column=1, padx=5, pady=5)
        
        underline_button = ttk.Button(format_window, text="Underline", command=lambda: self.apply_format("<u>"))
        underline_button.grid(row=0, column=2, padx=5, pady=5)
        
        heading_button = ttk.Button(format_window, text="Heading", command=self.add_heading)
        heading_button.grid(row=0, column=3, padx=5, pady=5)
        
        indent_button = ttk.Button(format_window, text="Indent", command=lambda: self.apply_format("<indent>"))
        indent_button.grid(row=1, column=0, padx=5, pady=5)
        
        bullet_button = ttk.Button(format_window, text="Bullet Point", command=lambda: self.apply_format("<bullet>"))
        bullet_button.grid(row=1, column=1, padx=5, pady=5)
        
        numbered_button = ttk.Button(format_window, text="Numbered List", command=lambda: self.apply_format("<numbered>"))
        numbered_button.grid(row=1, column=2, padx=5, pady=5)
        
    def apply_format(self, tag):
        if tag == "<indent>":
            self.textbox.insert(tk.INSERT, "\t")
        elif tag == "<bullet>":
            self.textbox.insert(tk.INSERT, "\n• ")
        elif tag == "<numbered>":
            self.textbox.insert(tk.INSERT, "\n1. ")
        else:
            selected_text = self.textbox.get("sel.first", "sel.last")
            formatted_text = f"{tag}{selected_text}{tag[::-1]}"
            self.textbox.delete("sel.first", "sel.last")
            self.textbox.insert("insert", formatted_text)
        
    def add_heading(self):
        heading_window = tk.Toplevel(self.root)
        heading_window.title("Heading")
        
        level_label = ttk.Label(heading_window, text="Heading Level (1-9):")
        level_label.grid(row=0, column=0, padx=5, pady=5)
        
        self.level_entry = ttk.Entry(heading_window)
        self.level_entry.grid(row=0, column=1, padx=5, pady=5)
        
        apply_button = ttk.Button(heading_window, text="Apply", command=self.apply_heading)
        apply_button.grid(row=1, columnspan=2, padx=5, pady=5)
        
    def apply_heading(self):
        level = self.level_entry.get()
        if level.isdigit() and 1 <= int(level) <= 9:
            selected_text = self.textbox.get("sel.first", "sel.last")
            formatted_text = f"<h{level}>{selected_text}</h{level}>"
            self.textbox.delete("sel.first", "sel.last")
            self.textbox.insert("insert", formatted_text)
        else:
            messagebox.showerror("Error", "Invalid heading level. Please enter a number between 1 and 9.")
        
    def save_document(self):
        content = self.textbox.get("1.0", "end-1c")
        filename = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if filename:
            document = Document()
            paragraph = document.add_paragraph()
            run = paragraph.add_run(content)
            run.font.size = Pt(12)
            document.save(filename)
            messagebox.showinfo("Success", f"Document '{filename}' saved successfully!")
        
    def open_document(self):
        filename = filedialog.askopenfilename(filetypes=[("Word Document", "*.docx")])
        if filename:
            document = Document(filename)
            content = ""
            for paragraph in document.paragraphs:
                content += paragraph.text + "\n"
            self.textbox.delete("1.0", "end")
            self.textbox.insert("1.0", content)
        
    def append_to_document(self):
        content = self.textbox.get("1.0", "end-1c")
        filename = filedialog.askopenfilename(filetypes=[("Word Document", "*.docx")])
        if filename:
            document = Document(filename)
            document.add_paragraph(content)
            document.save(filename)
            messagebox.showinfo("Success", "Content appended successfully!")

if __name__ == "__main__":
    root = tk.Tk()
    app = NotepadApp(root)
    root.mainloop()
